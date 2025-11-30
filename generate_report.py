"""Generate a .docx report summarizing the midterm project.

This script reads the previously generated plots and CSV results
and creates a Word document with short explanations under each figure,
as well as selected code blocks from the project.
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.config import Config, PathManager


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str):
    document.add_paragraph(text)


def add_figure(document: Document, image_path: Path, caption: str):
    if not image_path.exists():
        add_paragraph(document, f"[Görsel bulunamadı: {image_path.name}]")
        add_paragraph(document, caption)
        return
    document.add_picture(str(image_path), width=Inches(5.5))
    add_paragraph(document, caption)


def add_code_block(document: Document, code: str, title: str | None = None):
    """Add a monospace-styled code block to the document."""
    if title:
        add_paragraph(document, title)
    p = document.add_paragraph()
    run = p.add_run(code)
    # Set font to monospace (e.g., Consolas)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)


def read_file_snippet(path: Path, max_lines: int = 80) -> str:
    if not path.exists():
        return f"# Dosya bulunamadı: {path}"
    lines = path.read_text(encoding="utf-8").splitlines()
    if len(lines) > max_lines:
        lines = lines[:max_lines] + ["", "# ...dosya devam ediyor..."]
    return "\n".join(lines)


def main():
    config = Config()
    paths = PathManager(config)

    results_dir = config.results_dir
    plots_dir = config.plots_dir
    root_dir = Path(__file__).resolve().parent

    doc = Document()

    # Başlık
    doc.add_heading("Breast Cancer Wisconsin Sınıflandırma Projesi Raporu", level=0)

    # 1. Veri Seti ve Ön İşleme
    add_heading(doc, "1. Veri Seti ve Ön İşleme", level=1)
    add_paragraph(
        doc,
        "Bu projede scikit-learn kütüphanesinden Breast Cancer Wisconsin veri seti kullanılmıştır. "
        "Tüm özellikler sayısaldır ve veri setinde eksik değer bulunmamaktadır. Z-score yöntemi ile "
        "aykırı değerler analiz edilmiş, özellikle alan ve perimeter gibi değişkenlerde birkaç uç değer tespit edilmiştir.",
    )

    # İlgili kod blokları: data loader ve quality checker
    add_heading(doc, "1.1 Veri Yükleme ve Kalite Kontrolleri Kodları", level=2)
    loader_code = read_file_snippet(root_dir / "data" / "loader.py")
    add_code_block(doc, loader_code, "Kod 1: BreastCancerDataLoader")

    quality_code = read_file_snippet(root_dir / "data" / "quality.py")
    add_code_block(doc, quality_code, "Kod 2: DataQualityChecker")

    # 2. Keşifsel Veri Analizi
    add_heading(doc, "2. Keşifsel Veri Analizi (EDA)", level=1)

    # Korelasyon heatmap
    add_figure(
        doc,
        plots_dir / "correlation_heatmap.png",
        "Şekil 1: Özellikler arası Pearson korelasyon ısı haritası. Özellikle 'mean radius', 'mean perimeter' ve 'mean area' gibi boyutla ilgili değişkenler arasında yüksek pozitif korelasyon görülmektedir. Bu da bu değişkenlerin benzer bilgiyi taşıdığını ve PCA gibi yöntemlerle özetlenebileceğini gösterir.",
    )

    # Boxplot
    add_figure(
        doc,
        plots_dir / "boxplot_first10.png",
        "Şekil 2: İlk 10 özelliğin boxplot grafikleri. Özellikle 'mean area' ve 'mean perimeter' gibi değişkenlerde birkaç uç değer bulunmakta, ancak medyan etrafındaki dağılım genel olarak dengelidir. Aykırı değerler bazı lineer modelleri etkileyebilse de ağaç tabanlı yöntemler bu değerlere karşı daha dayanıklıdır.",
    )

    eda_code = read_file_snippet(root_dir / "eda" / "exploration.py")
    add_code_block(doc, eda_code, "Kod 3: EDAExplorer - tanımlayıcı istatistikler ve görselleştirmeler")

    # 3. Boyut İndirgeme
    add_heading(doc, "3. Boyut İndirgeme (PCA ve LDA)", level=1)

    # PCA explained variance
    add_figure(
        doc,
        plots_dir / "pca_explained_variance.png",
        "Şekil 3: PCA bileşenlerinin açıklanan varyans oranları. İlk birkaç bileşen toplam varyansın büyük kısmını açıklamakta ve bu sayede boyut indirgeme ile bilginin büyük kısmını korumak mümkündür.",
    )

    # PCA scatter
    add_figure(
        doc,
        plots_dir / "pca_2d_scatter.png",
        "Şekil 4: İlk iki PCA bileşeni (PC1 ve PC2) uzayında örneklerin dağılımı. İki sınıfın bu iki bileşen üzerinde belirgin şekilde ayrıştığı görülmekte, bu da orijinal özelliklerin yüksek ayrıştırma gücüne sahip olduğunu desteklemektedir.",
    )

    # LDA projection
    add_figure(
        doc,
        plots_dir / "lda_1d_hist.png",
        "Şekil 5: LDA ile elde edilen birinci diskriminant (LD1) boyunca sınıf dağılımları. İki sınıfın tek bir eksen üzerinde dahi belirgin şekilde ayrıştığı görülmektedir; bu da sınıflandırma probleminin lineer olarak da iyi ayrılabildiğini gösterir.",
    )

    pca_code = read_file_snippet(root_dir / "dimensionality_reduction" / "pca_module.py")
    add_code_block(doc, pca_code, "Kod 4: PCAReducer - PCA boyut indirgeme")

    lda_code = read_file_snippet(root_dir / "dimensionality_reduction" / "lda_module.py")
    add_code_block(doc, lda_code, "Kod 5: LDAReducer - LDA boyut indirgeme")

    # 4. Modelleme ve Değerlendirme
    add_heading(doc, "4. Modellerin Eğitimi ve Değerlendirilmesi", level=1)

    # Validation sonuç tablosu
    val_path = results_dir / "validation_results.csv"
    if val_path.exists():
        val_df = pd.read_csv(val_path)
        add_paragraph(
            doc,
            "Aşağıdaki tabloda, her üç veri temsili (ham, PCA, LDA) ve beş model kombinasyonu için validation seti üzerindeki performans metrikleri özetlenmiştir (accuracy, precision, recall, F1, ROC-AUC).",
        )
        table = doc.add_table(rows=1, cols=len(val_df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(val_df.columns):
            hdr_cells[i].text = col
        for _, row in val_df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(val_df.columns):
                row_cells[i].text = str(row[col])
    else:
        add_paragraph(doc, "validation_results.csv bulunamadı.")

    add_paragraph(
        doc,
        "Genel olarak logistic regression, ölçeklendirilmiş ham veri üzerinde en yüksek ROC-AUC değerine ulaşmış ve en iyi validation performansını göstermiştir. PCA ve LDA temsilleri de rekabetçi sonuçlar üretmiş, ancak ham veriye göre küçük kayıplar gözlenmiştir.",
    )

    # Confusion matrix
    add_figure(
        doc,
        plots_dir / "confusion_matrix.png",
        "Şekil 6: En iyi model (ham veri üzerinde Logistic Regression) için test seti confusion matrix'i. Doğru pozitif ve doğru negatif sayılarının yüksek olduğu, yanlış sınıflandırma sayısının ise oldukça düşük olduğu görülmektedir.",
    )

    # ROC curve
    add_figure(
        doc,
        plots_dir / "roc_curve.png",
        "Şekil 7: En iyi model için ROC eğrisi. Eğrinin sol üst köşeye yakın olması ve AUC değerinin 1'e çok yakın olması, modelin pozitif ve negatif sınıfları ayırmada çok başarılı olduğunu göstermektedir. Threshold değeri düşürüldüğünde true positive oranı artarken false positive oranı da artmaktadır; uygulama senaryosuna göre bu denge seçilmelidir.",
    )

    models_code = read_file_snippet(root_dir / "models" / "trainer.py")
    add_code_block(doc, models_code, "Kod 6: ModelTrainer - tüm modellerin eğitimi ve seçimi")

    registry_code = read_file_snippet(root_dir / "models" / "registry.py")
    add_code_block(doc, registry_code, "Kod 7: ModelRegistry - kullanılan modeller")

    # 5. SHAP ile Açıklanabilirlik Analizi
    add_heading(doc, "5. XAI – SHAP Analizi", level=1)

    add_paragraph(
        doc,
        "Bu bölümde en iyi modeller için SHAP (SHapley Additive exPlanations) yöntemi kullanılarak model kararlarının açıklanabilirliği incelenmiştir. Özellikle özelliklerin tahmin üzerindeki ortalama etkileri ve pozitif/negatif yöndeki katkıları analiz edilmiştir.",
    )

    # Raw (ham veri) için SHAP
    raw_bar = None
    raw_summary = None
    for p in plots_dir.glob("raw_*_shap_bar.png"):
        raw_bar = p
        break
    for p in plots_dir.glob("raw_*_shap_summary.png"):
        raw_summary = p
        break

    add_figure(
        doc,
        raw_bar if raw_bar else plots_dir / "best_raw_shap_bar.png",
        "Şekil 8: Ham veri üzerinde en iyi model için SHAP bar grafiği. En önemli birkaç özellik; tüm veri seti boyunca tahmin olasılığını en fazla değiştiren (ortalama mutlak SHAP değeri en yüksek olan) değişkenlerdir. Bu özellikler, tüm modellerde tekrarlayan şekilde öne çıkıyorsa, klinik açıdan da kritik değişkenler olarak yorumlanabilir.",
    )

    add_figure(
        doc,
        raw_summary if raw_summary else plots_dir / "best_raw_shap_summary.png",
        "Şekil 9: Ham veri üzerinde en iyi model için SHAP summary grafiği. Her nokta bir örneği temsil eder; renkler özelliğin değeri (düşük-yüksek), yatay eksen ise tahmin üzerindeki katkı yönü ve büyüklüğünü gösterir. Özelliğin yüksek değeri modelin malign (pozitif sınıf) tahmin etme olasılığını artırıyorsa noktalar sağ tarafa yoğunlaşmaktadır.",
    )

    # PCA SHAP
    pca_bar = None
    pca_summary = None
    for p in plots_dir.glob("pca_*_shap_bar.png"):
        pca_bar = p
        break
    for p in plots_dir.glob("pca_*_shap_summary.png"):
        pca_summary = p
        break

    add_figure(
        doc,
        pca_bar if pca_bar else plots_dir / "pca_rf_shap_bar.png",
        "Şekil 10: PCA ile indirgenmiş veri temsili için en iyi modelin SHAP bar grafiği. Burada önemli olanlar doğrudan orijinal özellikler değil, temel bileşenlerdir (PC1, PC2, vb.). Bu bileşenler, birden fazla orijinal özelliğin doğrusal kombinasyonları olduğundan, her bileşenin önemli olması; ilgili özellik grubunun modele güçlü katkı yaptığını gösterir.",
    )

    add_figure(
        doc,
        pca_summary if pca_summary else plots_dir / "pca_rf_shap_summary.png",
        "Şekil 11: PCA temsili için SHAP summary grafiği. Farklı PCA bileşenlerinin, gözlemler bazında model çıktısına nasıl katkı verdiği görülmektedir. Özellikle PC1 ve PC2 gibi bileşenlerin yüksek değeri, malign sınıf olasılığını artıran veya azaltan ana eksenleri temsil eder.",
    )

    # LDA SHAP
    lda_bar = None
    lda_summary = None
    for p in plots_dir.glob("lda_*_shap_bar.png"):
        lda_bar = p
        break
    for p in plots_dir.glob("lda_*_shap_summary.png"):
        lda_summary = p
        break

    add_figure(
        doc,
        lda_bar if lda_bar else plots_dir / "lda_best_shap_bar.png",
        "Şekil 12: LDA temsili için en iyi modelin SHAP bar grafiği. LDA bileşenleri (LD1, LD2, vb.), sınıflar arasındaki ayrımı maksimize edecek şekilde tanımlandığı için; burada en önemli bileşen genellikle LD1'dir. Bu bileşenin yüksek SHAP değeri, gözlemin ilgili sınıfa (örneğin malign) ait olma olasılığını güçlü şekilde etkiler.",
    )

    add_figure(
        doc,
        lda_summary if lda_summary else plots_dir / "lda_best_shap_summary.png",
        "Şekil 13: LDA temsili için SHAP summary grafiği. Sınıflar arasındaki ayrımı temsil eden diskriminant bileşenlerin, her bir gözlem için tahmin çıktısına katkısı görülmektedir. Özellikle LD1 eksenindeki değerlerin pozitif sınıf yönünde yoğunlaştığı gözlemler, model tarafından yüksek riskli (malign) olarak sınıflandırılmaktadır.",
    )

    shap_code = read_file_snippet(root_dir / "xai" / "shap_analysis.py")
    add_code_block(doc, shap_code, "Kod 8: SHAPAnalyzer - açıklanabilirlik analizleri")

    # 6. Pipeline Orkestrasyonu
    add_heading(doc, "6. Pipeline Orkestrasyonu", level=1)
    orchestrator_code = read_file_snippet(root_dir / "pipeline" / "orchestrator.py")
    add_code_block(doc, orchestrator_code, "Kod 9: PipelineOrchestrator - uçtan uca akış")

    add_paragraph(
        doc,
        "Sonuç olarak, Breast Cancer Wisconsin veri seti üzerinde kurulan klasik makine öğrenmesi modelleri oldukça yüksek performans göstermiştir. "
        "Özellikle logistic regression ve ağaç tabanlı yöntemler, uygun ölçeklendirme ve boyut indirgeme adımlarının ardından yüksek ROC-AUC ve F1 skorları üretmiştir. "
        "SHAP analizleri, modelin en çok hangi özelliklere dayandığını açıkça göstermiş, bu sayede model kararları klinik olarak yorumlanabilir hale gelmiştir.",
    )

    output_path = results_dir / "report_midterm.docx"
    doc.save(output_path)
    print(f"Rapor oluşturuldu: {output_path}")


if __name__ == "__main__":
    main()
